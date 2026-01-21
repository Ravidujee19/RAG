from __future__ import annotations

from typing import List

from docx import Document as DocxDocument

from app.models.document import Document


class DocxLoader:
    def load(self, source: str) -> List[Document]:
        doc = DocxDocument(source)
        text = "\n".join(p.text for p in doc.paragraphs if p.text)
        return [
            Document(
                id=f"{source}#docx",
                text=text,
                metadata={"source": source, "type": "docx"},
            )
        ]
